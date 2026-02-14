from scripts.paths import get_path
import subprocess
from pathlib import Path
import json
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil

from scripts.contracts import (
    OUTPUT_DIR,
    ACADEMIC_JSON_FILE,
    DESTINATION_DIR,
    TEMPLATES_DIR,
    DOCX_TEMPLATES_DIR,
    BLOOM_EXPANSION, CourseCategory
)

def print_table_styles(docx_path):
    """
    Prints all table styles available in the DOCX.
    Run once to discover the correct style name.
    """
    doc = Document(str(docx_path))

    print("\n📋 Available TABLE styles in this DOCX:\n")
    for style in doc.styles:
        if style.type == 3:  # WD_STYLE_TYPE.TABLE
            print(f"- {style.name!r}")

def force_table_style_design(
        docx_path: str | Path, 
        table_style: str = "Table Grid",
        header_row: bool = True,
        first_column: bool = True,
        banded_rows: bool = False,
        banded_columns: bool = False,
        total_row: bool = False,
        last_column: bool = False,
    ):
    doc = Document(str(docx_path))
    """
    Removes fixed table layout so Word re-applies AutoFit to Contents.
    """

    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            continue

        for el in tbl_pr.findall(qn("w:tblLayout")):
            tbl_pr.remove(el)
        
        table.style = table_style

        tbl_look = tbl_pr.find(qn("w:tblLook"))
        if tbl_look is None:
            tbl_look = OxmlElement("w:tblLook")
            tbl_pr.append(tbl_look)

        tbl_look.set(qn("w:firstRow"), "1" if header_row else "0")
        tbl_look.set(qn("w:firstColumn"), "1" if first_column else "0")
        tbl_look.set(qn("w:lastRow"), "1" if total_row else "0")
        tbl_look.set(qn("w:lastColumn"), "1" if last_column else "0")
        tbl_look.set(qn("w:noHBand"), "0" if banded_rows else "1")
        tbl_look.set(qn("w:noVBand"), "0" if banded_columns else "1")

    doc.save(str(docx_path))



def generate_word_co_bloom():
    json_path = get_path(OUTPUT_DIR, ACADEMIC_JSON_FILE)
    if not json_path.exists():
        raise FileNotFoundError(
            f"Stage 8 DOCX: Academic JSON not found at {json_path}"
        )
    """Generates a 5-column Word table with AutoFit to Contents behavior."""
    md_path = get_path(OUTPUT_DIR, "temp_report.md")
    docx_path = get_path(DESTINATION_DIR, "CO_BLOOM_Table.docx")
    ref_doc_path = get_path(TEMPLATES_DIR, DOCX_TEMPLATES_DIR, "co-bloom-reference.docx")

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    if "courses" not in data:
        raise ValueError("Stage 8 DOCX: Missing 'courses' key in academic JSON.")
    
    courses_by_category = data["courses"]

    if not isinstance(courses_by_category, dict):
        raise TypeError("Stage 8 DOCX: 'courses' must be a dictionary.")

    if not courses_by_category:
        raise RuntimeError("Stage 8 DOCX: 'courses' dictionary is empty.")

    valid_categories = {cat.code for cat in CourseCategory}

    total_count = 0

    for category, courses in courses_by_category.items():

        if category not in valid_categories:
            raise ValueError(
                f"Stage 8 DOCX: Invalid category '{category}'."
            )

        if not isinstance(courses, list):
            raise TypeError(
                f"Stage 8 DOCX: Category '{category}' must contain a list."
            )

        total_count += len(courses)

    if total_count == 0:
        raise RuntimeError(
            "Stage 8 DOCX: Categories exist but contain zero courses."
        )

    with open(md_path, "w", encoding="utf-8") as f:
        
        # Using a simple pipe table. 
        # Note: Do not use excessive dashes, as Pandoc uses them to calculate fixed widths.
        f.write("| CO-ID | Course Outcome | Course Code | Bloom's Level | Justification / Remarks |\n")
        f.write("|:---|:---|:---|:---|:---|\n")
        
        for cat in CourseCategory:
            courses = courses_by_category.get(cat.code, [])

            for course in courses:
                code = course.get("course_code")

                if not code:
                    raise ValueError(
                        f"Stage 8 DOCX: Course missing 'course_code' in category '{cat.code}'."
                    )

                for idx, co in enumerate(course.get("outcomes", []), start=1):

                    co_id = f"{code}.{idx}"
                    text = co.get("outcome", "").replace("\n", " ")

                    k_lvl = co.get("k_level")
                    bloom_code = co.get("bloom")
                    bloom_full = BLOOM_EXPANSION.get(bloom_code, "Unknown")

                    bloom_display = f"K{k_lvl} ({bloom_code})" if k_lvl else "N/A"

                    f.write(f"| {co_id} | {text} | {code} | {bloom_display} | |\n")

    if not shutil.which("pandoc"):
        raise RuntimeError("Stage 8 DOCX: Pandoc not found in system path.")


    try:
        # The key is using the reference-doc which has the 'Table' style set to AutoFit
        subprocess.run([
            "pandoc", 
            "-s",
            str(md_path), 
            "-o", str(docx_path),
            f"--reference-doc={str(ref_doc_path)}"
        ], check=True)

        force_table_style_design(docx_path, "Grid Table 6 Colorful")
        print(f"✅ Word Document Generated: {docx_path}")

    except Exception as e:
        raise RuntimeError(f"Stage 8 DOCX: Conversion failed: {e}")
    finally:
        if md_path.exists():
            md_path.unlink()