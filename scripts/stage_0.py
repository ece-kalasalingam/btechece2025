from pathlib import Path
from typing import List, Optional

from scripts.contracts import COURSES_DIR, INDEX_FILE
from scripts.paths import get_path
from scripts.utils import normalize_line_endings, validate_course_code

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None


def _style_to_markdown_prefix(style_name: str) -> str:
    """Map common DOCX paragraph styles to markdown-like prefixes."""
    normalized = (style_name or "").strip().lower()

    if normalized.startswith("heading"):
        # Handles names like "Heading 1", "Heading 2", etc.
        parts = normalized.split()
        level = 1
        if len(parts) > 1 and parts[1].isdigit():
            level = int(parts[1])
        level = max(1, min(level, 6))
        return "#" * level + " "

    if "list bullet" in normalized:
        return "- "

    if "list number" in normalized:
        # Keep markdown syntax compatible with existing downstream checks.
        return "1. "

    return ""


def _table_to_markdown(table) -> str:
    """Convert a DOCX table to a markdown table block."""
    rows = []
    for row in table.rows:
        cells = [normalize_line_endings(cell.text).strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    header = rows[0]
    body = rows[1:] if len(rows) > 1 else []

    markdown_lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * len(header)) + " |",
    ]

    for row in body:
        padded = row + [""] * (len(header) - len(row))
        markdown_lines.append("| " + " | ".join(padded[:len(header)]) + " |")

    return "\n".join(markdown_lines)


def docx_to_markdown_text(docx_path: Path) -> str:
    """Extract text from DOCX in a markdown-like structure expected by downstream stages."""
    if Document is None:
        raise RuntimeError(
            "python-docx is required to ingest .docx course files. "
            "Install dependencies from requirements.txt."
        )

    doc = Document(str(docx_path))
    lines: List[str] = []

    paragraph_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)

    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = next(paragraph_iter)
            text = normalize_line_endings(paragraph.text).strip()
            if not text:
                continue
            prefix = _style_to_markdown_prefix(getattr(paragraph.style, "name", ""))
            lines.append(f"{prefix}{text}")
        elif child.tag.endswith("}tbl"):
            table = next(table_iter)
            table_block = _table_to_markdown(table)
            if table_block:
                if lines and lines[-1] != "":
                    lines.append("")
                lines.extend(table_block.splitlines())
                lines.append("")

    return "\n".join(lines).strip() + "\n"


def load_course_source(file_path_md: Path, file_path_docx: Path) -> tuple[Optional[str], Optional[str]]:
    """
    Load course content, preferring DOCX over Markdown.
    Returns (raw_text, error_message).
    """
    if file_path_docx.exists():
        try:
            return docx_to_markdown_text(file_path_docx), None
        except Exception as exc:
            return None, f"Failed to parse DOCX file {file_path_docx}: {exc}"

    if file_path_md.exists():
        try:
            return file_path_md.read_text(encoding="utf-8"), None
        except Exception as exc:
            return None, f"Failed to read Markdown file {file_path_md}: {exc}"

    return None, (
        "Missing course file. Expected one of: "
        f"{file_path_docx.name} or {file_path_md.name}"
    )

def ingest(raw_text: str) -> str:
    """
    Pure text preparation. 
    Standardizes line endings for cross-platform (Windows/Linux/GitHub) compatibility.
    """
    return normalize_line_endings(raw_text)

def parse_index(index_path: Path) -> List[str]:
    """
    Parses index.md and returns a list of course codes.
    Ignores blank lines and comments.
    """
    if not index_path.exists():
        raise FileNotFoundError(f"Index file not found at {index_path}")

    course_codes = []
    with index_path.open(encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            if line.startswith("#"):
                continue
            if line.startswith(("-", "*", "+")):
                line = line[1:].strip()
            course_codes.append(line)

    return course_codes

def iter_courses():
    """
    Yields (course_code, raw_text) one at a time.
    """
    courses_path = get_path(COURSES_DIR)
    index_path = courses_path / INDEX_FILE

    course_codes = parse_index(index_path)
    course_codes = sorted(course_codes)

    for code in course_codes:
        try:
            validate_course_code(code)
        except ValueError as e:
            yield code, None, str(e)
            continue

        file_path_md = courses_path / f"{code}.md"
        file_path_docx = courses_path / f"{code}.docx"

        raw_text, error = load_course_source(file_path_md, file_path_docx)
        if error is not None:
            yield code, None, error
            continue

        assert raw_text is not None
        yield code, ingest(raw_text), None
